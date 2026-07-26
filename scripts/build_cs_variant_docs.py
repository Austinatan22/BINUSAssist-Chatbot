# -*- coding: utf-8 -*-
"""Scrape the per-campus / online / graduate Computer Science program cards from
binus.ac.id/program/ and assemble one program document per variant in documents_dir, so
get_program_catalog picks each up as a routable program distinct from the flagship
Computer Science / Data Science catalogs (KB Expansion Task 5, scope = "new programs only").

The six flagship CS-family programs (CS, AI, Data Science, Cyber Security, Software
Engineering, Global Class) already have their own richer curriculum PDFs, so this only adds
the eight programs the KB genuinely lacked: four per-campus CS variants, two BINUS Online
variants, and the Master / Doctor of Computer Science graduate programs.

Each doc opens with a disambiguation paragraph (same lesson as CS International in Task 4):
without it, "Computer Science Medan" collides with plain "Computer Science" under the
span-based program-name absorption in generation.py. Reindex-safe: these are real local
files like every other catalog entry, so /admin/reindex re-ingests them offline with no
BINUS dependency.

The two Master URLs on binus.ac.id/program/ (master-of-comp-science + master-of-
information-technology) serve BYTE-IDENTICAL "Master of Computer Science" content, so only
one is ingested here.
"""
import re
import urllib.request

import trafilatura
from docx import Document

from backend.config import settings


def get(u: str) -> str:
    req = urllib.request.Request(u, headers={"User-Agent": "Mozilla/5.0"})
    return urllib.request.urlopen(req, timeout=25).read().decode("utf-8", "ignore")


def clean_text(u: str) -> str:
    return trafilatura.extract(get(u), url=u, include_tables=True) or ""


# (filename stem, heading, disambiguation paragraph, source URL)
PROGRAMS = [
    (
        "Computer_Science_Medan_2026",
        "Computer Science (@Medan) - BINUS University",
        "This document describes the Computer Science program offered at the BINUS @Medan "
        "campus. It is the Medan-campus variant of Computer Science and is distinct from the "
        "main Computer Science program, from Computer Science Global Class, and from Computer "
        "Science International.",
        "https://binus.ac.id/program/computer-science-binus-medan/",
    ),
    (
        "Computer_Science_Semarang_2026",
        "Computer Science (@Semarang) - BINUS University",
        "This document describes the Computer Science program offered at the BINUS @Semarang "
        "campus. It is the Semarang-campus variant of Computer Science and is distinct from "
        "the main Computer Science program, from Computer Science Global Class, and from "
        "Computer Science International.",
        "https://binus.ac.id/program/computer-science-semarang/",
    ),
    (
        "Computer_Science_Malang_2026",
        "Computer Science (@Malang) - BINUS University",
        "This document describes the Computer Science program offered at the BINUS @Malang "
        "campus. It is the Malang-campus variant of Computer Science and is distinct from the "
        "main Computer Science program, from Computer Science Global Class, and from Computer "
        "Science International.",
        "https://binus.ac.id/program/computer-science-malang/",
    ),
    (
        "Computer_Science_Bandung_2026",
        "Computer Science (@Bandung) - BINUS University",
        "This document describes the Computer Science program offered at the BINUS @Bandung "
        "campus. It is the Bandung-campus variant of Computer Science and is distinct from the "
        "main Computer Science program, from Computer Science Global Class, and from Computer "
        "Science International.",
        "https://binus.ac.id/program/computer-science-bandung/",
    ),
    (
        "Computer_Science_Online_2026",
        "Computer Science (BINUS Online) - BINUS University",
        "This document describes the Computer Science program offered through BINUS Online "
        "Learning (online / distance mode for working professionals). It is distinct from the "
        "on-campus Computer Science program, from Computer Science Global Class, and from "
        "Computer Science International.",
        "https://binus.ac.id/program/computer-science-3/",
    ),
    (
        "Data_Science_Online_2026",
        "Data Science (BINUS Online) - BINUS University",
        "This document describes the Data Science program offered through BINUS Online "
        "Learning (online / distance mode for working professionals). It is distinct from the "
        "on-campus Data Science program.",
        "https://binus.ac.id/program/computer-science-data-science/",
    ),
    (
        "Master_of_Computer_Science_2026",
        "Master of Computer Science - BINUS Graduate Program",
        "This document describes the Master of Computer Science, a graduate (S2 / Master's) "
        "program under the BINUS Graduate Program, School of Computer Science. It is a "
        "postgraduate degree, distinct from the undergraduate (S1) Computer Science program "
        "and from the Doctor of Computer Science.",
        "https://binus.ac.id/program/master-of-comp-science/",
    ),
    (
        "Doctor_of_Computer_Science_2026",
        "Doctor of Computer Science - BINUS Graduate Program",
        "This document describes the Doctor of Computer Science, a doctoral (S3 / PhD) program "
        "under the BINUS Graduate Program, School of Computer Science. It is a doctoral "
        "degree, distinct from the undergraduate (S1) Computer Science program and from the "
        "Master of Computer Science.",
        "https://binus.ac.id/program/doctor-of-computer-science/",
    ),
]


def main() -> None:
    for stem, heading, disambig, url in PROGRAMS:
        text = clean_text(url)
        doc = Document()
        doc.add_heading(heading, level=0)
        doc.add_paragraph(disambig)
        doc.add_heading("Program Overview", level=1)
        for para in re.split(r"\n{1,}", text):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
        out = settings.documents_dir / f"{stem}.docx"
        doc.save(out)
        print(f"wrote {out.name}  ({len(text)} chars scraped)")


if __name__ == "__main__":
    main()
