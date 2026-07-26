# -*- coding: utf-8 -*-
"""Scrape the BINUS International Computer Science pages and assemble a program document
(Computer_Science_International_2026.docx) in documents_dir, so get_program_catalog picks it
up as the routable program "Computer Science International" (distinct from Computer Science /
Computer Science Global Class). Reindex-safe: it's a real local file like the other catalogs."""
import re, urllib.request
import trafilatura
from docx import Document
from backend.config import settings

def get(u):
    return urllib.request.urlopen(urllib.request.Request(u, headers={'User-Agent': 'Mozilla/5.0'}), timeout=25).read().decode('utf-8', 'ignore')

def clean_text(u):
    return trafilatura.extract(get(u), url=u, include_tables=True) or ''

SECTIONS = [
    ("About the Program", "https://international.binus.ac.id/computer-science/about-us/"),
    ("Course Structure", "https://international.binus.ac.id/computer-science/academic-programs/course-structure/"),
    ("Career Outlook", "https://international.binus.ac.id/computer-science/academic-programs/career-outlook/"),
]

doc = Document()
doc.add_heading("Computer Science International - BINUS International", level=0)
doc.add_paragraph(
    "This document describes the Computer Science program offered by BINUS International "
    "(referred to here as Computer Science International). It is a distinct program from the "
    "regular Computer Science program and from Computer Science Global Class. The BINUS "
    "International Computer Science program is taught fully in English and offers international "
    "dual-degree pathways (e.g. with La Trobe University)."
)

total = 0
for heading, url in SECTIONS:
    text = clean_text(url)
    total += len(text)
    doc.add_heading(heading, level=1)
    for para in re.split(r'\n{1,}', text):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    print(f"  {heading}: {len(text)} chars")

out = settings.documents_dir / "Computer_Science_International_2026.docx"
doc.save(out)
print(f"\nwrote {out}  (total scraped {total} chars)")
